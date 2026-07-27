from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT = DOCS / "TimeTracker_User_Manual_Training_2026-05-16.docx"
LOGO = ROOT / "public" / "logo.jpg"


ACCENT = "0EA5E9"
NAVY = "0F172A"
MUTED = "64748B"
LIGHT = "EFF6FF"
GREEN = "DCFCE7"
AMBER = "FEF3C7"
ROSE = "FFE4E6"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Prompt"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Prompt")
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="CBD5E1") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Prompt"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Prompt")
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(4)

    for name, size, color in [
        ("Title", 25, NAVY),
        ("Heading 1", 17, ACCENT),
        ("Heading 2", 13, NAVY),
        ("Heading 3", 11.5, NAVY),
    ]:
        style = doc.styles[name]
        style.font.name = "Prompt"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Prompt")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def add_cover(doc: Document) -> None:
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.15))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("คู่มือการใช้งาน TimeTracker")
    r.font.name = "Prompt"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Prompt")
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("สำหรับพนักงาน หัวหน้างาน HR และผู้ดูแลระบบ")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    meta.columns[0].width = Cm(4)
    meta.columns[1].width = Cm(10)
    rows = [
        ("เวอร์ชันเอกสาร", "Training Edition 2026-05-16"),
        ("ระบบ", "TimeTracker Web App"),
        ("กลุ่มผู้ใช้งาน", "Employee, Manager, HR, Admin"),
        ("โซนเวลา", "Asia/Bangkok (UTC+7)"),
    ]
    for i, row in enumerate(rows):
        set_cell_text(meta.cell(i, 0), row[0], bold=True, color=NAVY)
        set_cell_text(meta.cell(i, 1), row[1])
        set_cell_shading(meta.cell(i, 0), LIGHT)
    set_table_borders(meta)

    box = doc.add_table(rows=1, cols=1)
    set_cell_shading(box.cell(0, 0), LIGHT)
    set_cell_text(
        box.cell(0, 0),
        "จุดประสงค์: ใช้เป็นเอกสารประกอบการเทรนนิ่งและเป็นคู่มืออ้างอิงหลังเริ่มใช้งานจริง ครอบคลุมวิธีใช้งาน กติกาการบันทึกเวลา การลา OT On-site รายงาน และงานผู้ดูแลระบบ",
        bold=False,
    )
    set_table_borders(box, color="BFDBFE")
    doc.add_page_break()


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def p(doc, text, bold_prefix: str | None = None):
    para = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = para.add_run(bold_prefix)
        run.bold = True
        para.add_run(text[len(bold_prefix):])
    else:
        para.add_run(text)
    return para


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def callout(doc, title: str, body: str, fill: str = LIGHT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    para = cell.paragraphs[0]
    r = para.add_run(title + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    para.add_run(body)
    set_table_borders(table, color="CBD5E1")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    if widths:
        for i, w in enumerate(widths):
            t.columns[i].width = Cm(w)
    for i, header in enumerate(headers):
        set_cell_text(t.cell(0, i), header, bold=True, color="FFFFFF")
        set_cell_shading(t.cell(0, i), ACCENT)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    set_table_borders(t)
    return t


def add_footer(doc: Document) -> None:
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("TimeTracker User Manual | Training Edition | 2026-05-16")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def build_manual() -> None:
    DOCS.mkdir(exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_cover(doc)

    h1(doc, "1. ภาพรวมระบบ")
    p(doc, "TimeTracker เป็น Web App สำหรับบันทึกเวลาเข้า-ออกงาน, งานนอกสถานที่, OT, ใบลา, Daily Report และการตรวจสอบข้อมูลเวลาทำงานของบริษัท โดยข้อมูลหลักเชื่อมกับบัญชีผู้ใช้และสิทธิ์ตามบทบาท")
    table(
        doc,
        ["บทบาท", "สิทธิ์หลัก", "เมนูที่เกี่ยวข้อง"],
        [
            ("Employee", "บันทึกเวลา, ส่ง Daily Report, ยื่น OT/ใบลา, ดูประวัติของตนเอง", "Home, On-site, Requests, Report, Calendar, Profile"),
            ("Manager", "อนุมัติ/ไม่อนุมัติคำขอของทีม, ดูคำขอที่รออนุมัติ", "Requests, Profile"),
            ("HR/Admin", "ดูสรุปพนักงาน, แก้ไขเวลา, จัดการสิทธิ์, วันหยุด, นโยบายลา และ export รายงาน", "HR, Audit, Team, Time Sync, Settings, QR Display"),
            ("Viewer/บัญชีรออนุมัติ", "รอ Admin เปิดสิทธิ์ก่อนใช้งาน", "Pending Approval / Access Suspended"),
        ],
        widths=[3, 8, 5],
    )
    callout(doc, "หลักการสำคัญ", "ข้อมูลเวลาและคำขอในระบบถือเป็นข้อมูลอ้างอิงการทำงาน ผู้ใช้ควรกดบันทึกด้วยตนเองและตรวจสอบความถูกต้องทุกวัน", AMBER)

    h1(doc, "2. การเข้าสู่ระบบและเมนู")
    h2(doc, "2.1 เข้าสู่ระบบ")
    numbered(doc, [
        "เปิด Web App TimeTracker จากลิงก์ของบริษัท",
        "กด Sign in with Google และเลือกบัญชีบริษัท",
        "หากเป็นผู้ใช้ใหม่ สถานะอาจเป็นรออนุมัติ ต้องให้ Admin เปิดสิทธิ์เป็น active ก่อน",
        "หากบัญชีถูกระงับ ระบบจะแสดงหน้าแจ้งเตือนและไม่สามารถใช้งานต่อได้",
    ])
    h2(doc, "2.2 เมนูหลัก")
    table(
        doc,
        ["เมนู", "ใช้ทำอะไร"],
        [
            ("Home / Time Attendance", "เช็คอินโรงงาน, เช็คเอาท์, เริ่ม/จบ OT, ดูสรุปวันนี้และกราฟรายสัปดาห์"),
            ("On-site", "สร้างหรือเข้าห้องงานนอกสถานที่, Check-in/Check-out กลุ่ม, กลับโรงงาน, ออกก่อน"),
            ("Requests", "ยื่นคำขอ OT และใบลา, ตรวจสถานะ, ยกเลิกใบลาที่อนุมัติแล้ว, อนุมัติคำขอสำหรับ Manager"),
            ("Report", "บันทึก Daily Report รายวัน"),
            ("Calendar", "ดูปฏิทินการทำงาน วันลา วันหยุด และสถานะรายวัน"),
            ("Profile", "ดูสรุปของตนเอง เช่น มาสาย, OT, อัตราส่งรายงาน และประวัติรายวัน"),
            ("HR / Audit / Settings", "เมนูผู้ดูแล ใช้ตรวจสอบ แก้ไข และตั้งค่าระบบ"),
        ],
        widths=[4, 12],
    )

    h1(doc, "3. กติกาเวลาทำงาน")
    table(
        doc,
        ["หัวข้อ", "กฎที่ระบบใช้"],
        [
            ("เวลาเริ่มงาน", "08:30 น. หาก Check-in หลังเวลา threshold จะถือว่า late ยกเว้นวันหยุดหรือมีใบลาที่อนุมัติแล้ว"),
            ("เวลางานปกติ", "ระบบคำนวณช่วง 08:30-17:30 และหักพักกลางวัน 12:00-13:00"),
            ("Auto Check-out วันทำงาน", "ถ้า Check-in แล้วไม่ได้ Check-out ระบบจะปิดเวลาให้อัตโนมัติที่ 17:30 น."),
            ("วันหยุด/เสาร์-อาทิตย์", "ไม่นับสายเมื่อมีการทำงานวันหยุด และมีเงื่อนไขสะสมสิทธิ์แลกวันหยุด"),
            ("Safety Check-out วันหยุด", "ถ้าวันหยุดยังไม่ได้ Check-out ระบบ safety checkout เวลา 23:00 น."),
            ("OT หน้า Home", "ปุ่ม Start OT เปิดใช้งานตั้งแต่ 18:00 น. เป็นต้นไป"),
            ("OT On-site", "นับ OT จาก 17:30 น. เป็นต้นไป และสามารถหักเวลาพักช่วง OT ตอนปิด session"),
        ],
        widths=[4, 12],
    )
    h2(doc, "3.1 ใบลามีผลต่อการนับสาย")
    table(
        doc,
        ["ประเภทการลา", "ผลต่อเวลาเข้างาน"],
        [
            ("ลาทั้งวัน", "สถานะเป็น leave ไม่ต้อง Check-in"),
            ("ลาครึ่งเช้า", "เข้างานช่วงบ่ายได้ถึง 13:00 น. หลังจากนั้นถือว่าสาย"),
            ("ลาครึ่งบ่าย", "เวลาเข้างานยังใช้ 08:30 น. ตามปกติ"),
            ("ลารายชั่วโมงตั้งแต่ต้นวัน", "threshold เลื่อนไปตามเวลาสิ้นสุดการลา เช่น ลา 08:30-10:30 ให้เข้างานได้ถึง 10:30"),
            ("ลารายชั่วโมงช่วงอื่น", "ไม่กระทบเวลาเข้างานเช้า ใช้ 08:30 น."),
        ],
        widths=[5, 11],
    )

    h1(doc, "4. การ Check-in / Check-out โรงงาน")
    h2(doc, "4.1 Check-in ด้วย QR")
    numbered(doc, [
        "อยู่บริเวณโรงงานและเปิดหน้า Home",
        "กดสแกน QR หรือเปิดกล้องตามที่ระบบถามสิทธิ์",
        "สแกน QR จากจอ QR Display ของบริษัท",
        "เมื่อสำเร็จ ระบบบันทึก first_check_in ของวันนั้น และแสดงเวลาบน Daily Summary",
    ])
    callout(doc, "ข้อควรรู้เกี่ยวกับ QR", "QR มีอายุสั้นและมี nonce ป้องกันการใช้ซ้ำ หากขึ้นว่าหมดอายุหรือถูกใช้แล้ว ให้สแกน QR ใหม่จากจอ", LIGHT)
    h2(doc, "4.2 Check-out")
    numbered(doc, [
        "กด Check-out ที่หน้า Home เมื่อเลิกงาน",
        "ระบบจะแสดงหน้าต่างยืนยันและเตือนเรื่อง Auto Check-out 17:30",
        "หากยังมี On-site Session ค้างอยู่ ระบบจะแจ้งเตือนให้ Check-out ที่หน้า On-site ก่อน",
        "หลัง Check-out ระบบสรุปเวลาทำงานปกติและ OT ที่เกี่ยวข้อง",
    ])
    callout(doc, "กรณีลืม Check-out", "วันทำงานปกติระบบจะ Auto Check-out เวลา 17:30 และบันทึก regular_hours = 8 ชั่วโมง แต่ควร Check-out เองทุกครั้งเพื่อให้ timeline ถูกต้อง", AMBER)

    h1(doc, "5. การทำ OT")
    h2(doc, "5.1 OT ที่บันทึกจากหน้า Home")
    numbered(doc, [
        "ต้อง Check-in และเลิกงานปกติก่อน จึงเข้าสู่ช่วง OT",
        "ปุ่ม Start OT เปิดเวลา 18:00 น.",
        "กด Start OT เมื่อเริ่มทำ OT จริง",
        "กด End OT เมื่อจบงาน OT",
        "ระบบแสดง OT elapsed และสรุปชั่วโมง OT ใน Daily Summary",
    ])
    h2(doc, "5.2 การยื่น OT Request")
    numbered(doc, [
        "ไปที่ Requests แล้วกดปุ่ม OT Request",
        "เลือกวันที่, เวลาเริ่ม, เวลาสิ้นสุด, โครงการ และระบุเหตุผล",
        "ตรวจสอบจำนวนชั่วโมงที่ระบบคำนวณ",
        "กดยื่นคำขอ สถานะจะเป็นรออนุมัติ",
        "Manager/Admin อนุมัติหรือไม่อนุมัติได้จากแท็บรออนุมัติ",
    ])
    table(
        doc,
        ["สถานะ", "ความหมาย"],
        [
            ("pending", "รออนุมัติ"),
            ("approved", "อนุมัติแล้ว นำไปแสดงในสรุป OT"),
            ("rejected", "ไม่อนุมัติ พร้อมเหตุผลถ้ามี"),
        ],
        widths=[4, 12],
    )

    h1(doc, "6. ใบลาและกฎวันลา")
    h2(doc, "6.1 ประเภทใบลา")
    table(
        doc,
        ["ประเภท", "คำอธิบาย", "รูปแบบที่รองรับ"],
        [
            ("ลาป่วย", "เจ็บป่วย/พบแพทย์", "ทั้งวัน, ครึ่งวัน, รายชั่วโมง"),
            ("ลากิจ", "ธุระส่วนตัวหรือเหตุจำเป็น", "ทั้งวัน, ครึ่งวัน, รายชั่วโมง"),
            ("ลาพักร้อน", "วันหยุดพักผ่อนประจำปี", "ทั้งวัน, ครึ่งวัน, รายชั่วโมงตามนโยบายระบบ"),
            ("ลากิจพิเศษ", "เหตุพิเศษตามนโยบายบริษัท", "ทั้งวัน, ครึ่งวัน, รายชั่วโมง"),
            ("แลกวันหยุด", "ใช้สิทธิ์ที่ได้จากการทำงานวันหยุด", "ใช้ตามยอดสิทธิ์คงเหลือ"),
            ("ลาอื่น ๆ", "ระบุเหตุผลเพิ่มเติม", "ตามที่ผู้จัดการพิจารณา"),
        ],
        widths=[4, 6, 6],
    )
    h2(doc, "6.2 วิธีขอใบลา")
    numbered(doc, [
        "ไปที่ Requests แล้วเลือก Leave Request",
        "เลือกประเภทการลา",
        "เลือกช่วงเวลา: ทั้งวัน, ครึ่งเช้า, ครึ่งบ่าย หรือรายชั่วโมง",
        "เลือกวันที่เริ่มต้นและวันที่สิ้นสุด",
        "กรอกเหตุผลให้ชัดเจน",
        "กดยื่นใบลาและรอการอนุมัติ",
    ])
    callout(doc, "ลาพักร้อน", "ระบบแจ้งเตือนว่าลาพักร้อนต้องยื่นล่วงหน้าอย่างน้อย 3 วัน โดยนับรวมวันที่ทำรายการ", AMBER)
    h2(doc, "6.3 สิทธิ์ลาและการตัดยอด")
    bullets(doc, [
        "ยอดสิทธิ์วันลาจะตัดเมื่อใบลาถูกอนุมัติ",
        "ถ้าขอเกินสิทธิ์คงเหลือ ระบบยังยื่นได้ แต่จะแสดงคำเตือนและต้องรอหัวหน้าอนุมัติ",
        "ใบลาที่ถูกปฏิเสธหรือยกเลิกสำเร็จจะคืนผลต่อการนับสถานะเวลา",
        "ใบลาที่อนุมัติแล้วจะมีปุ่ม ยกเลิกใบลา เมื่อกดยืนยัน ระบบจะเปลี่ยนสถานะเป็น cancelled ทันที",
        "เมื่อยกเลิกใบลา ระบบคืน balance ทันที และเรียก DB RPC เพื่อล้างหรือคำนวณ daily_time_logs ใหม่ทันที",
        "ระบบยังรองรับข้อมูลเดิมที่ค้างสถานะ cancel_requested เพื่อให้ Manager/Admin ตรวจและปิดเคสเก่าได้",
    ])
    h2(doc, "6.4 การยกเลิกใบลาที่อนุมัติแล้ว")
    numbered(doc, [
        "ไปที่ Requests แล้วเปิดแท็บ Leave Request",
        "เลือกใบลาที่สถานะอนุมัติแล้ว",
        "กดปุ่ม ยกเลิกใบลา",
        "ตรวจสอบรายละเอียดและกดยืนยัน",
        "ระบบเปลี่ยนสถานะใบลาเป็น cancelled ทันที คืน balance ทันที และปรับ daily_time_logs ผ่าน DB RPC ทันที",
    ])
    callout(doc, "ข้อมูลเก่าที่รอยกเลิก", "ถ้าเคยมีรายการเก่าที่อยู่สถานะ cancel_requested ระบบยังรองรับอยู่เพื่อให้ Manager/Admin จัดการข้อมูลค้างเดิมได้ แต่ flow ใหม่ของพนักงานจะไม่เข้า รอยกเลิก แล้ว", AMBER)

    h1(doc, "7. Daily Report")
    p(doc, "Daily Report ใช้บันทึกว่าวันนั้นทำงานให้ลูกค้า/โครงการใด รายละเอียดงานอะไร และช่วงเวลาใด ระบบรองรับหลายรายการต่อวัน")
    numbered(doc, [
        "ไปที่ Report หรือกด popup วางแผนงานจากหน้า Home",
        "เลือกวันที่รายงาน",
        "เลือก Detail ของงาน",
        "เลือก End User และ Project No. หากไม่มีในรายการให้เลือก Other แล้วกรอกเอง",
        "เลือก Working Period: ALL, HALF DAY เช้า, HALF DAY บ่าย หรือ SOME TIME",
        "กดเพิ่มรายการงานหากมีหลายงานในวันเดียวกัน",
        "กดบันทึก Daily Report เมื่อกรอกครบ",
    ])
    callout(doc, "การลบ/แก้ไข", "รายการที่บันทึกแล้วจะขึ้นสถานะบันทึกแล้ว หากลบรายการสุดท้ายของวัน ระบบจะลบหัวรายงานของวันนั้นด้วย", LIGHT)

    h1(doc, "8. On-site")
    h2(doc, "8.1 ภาพรวม")
    p(doc, "On-site ใช้สำหรับงานนอกสถานที่แบบเป็นกลุ่ม โดยมี Leader เป็นผู้สร้างห้อง เลือกโครงการและสมาชิก แล้วกด Check-in/Check-out ให้กลุ่ม")
    h2(doc, "8.2 สร้างห้อง On-site")
    numbered(doc, [
        "ไปที่ On-site แล้วกด สร้างห้อง",
        "เลือก End User และ Project No. หรือเลือก Other แล้วกรอกชื่อ End User",
        "เลือกสมาชิกทีม ระบบนับรวมผู้สร้างเป็น Leader โดยอัตโนมัติ",
        "กดสร้างห้อง On-site",
        "เข้าห้องที่สร้างเพื่อกด Check-in เมื่อเริ่มงานจริง",
    ])
    h2(doc, "8.3 ระหว่างทำ On-site")
    bullets(doc, [
        "สถานะ open หมายถึงสร้างห้องแล้วแต่ยังไม่ได้ Check-in",
        "สถานะ checked_in หมายถึงกำลังทำงานนอกสถานที่",
        "Leader สามารถเพิ่มสมาชิกระหว่างวันได้ สมาชิกที่เพิ่มกลางวันจะใช้เวลาที่เพิ่มจริง และไม่ได้เบี้ยเลี้ยงเช้า",
        "สมาชิกที่ออกก่อนให้ใช้ฟังก์ชัน Early Leave พร้อมเหตุผล ระบบจะบันทึกเวลาออกของคนนั้น",
        "ถ้ากลับโรงงานก่อน 17:30 ให้เลือกกลับโรงงาน ระบบเปลี่ยน work_type เป็น mixed และรอ Auto Check-out 17:30",
        "เมื่อปิดงานทั้งกลุ่มหลัง 17:30 ระบบถามเวลาพักช่วง OT เพื่อนำไปหักจาก OT จริง",
    ])
    h2(doc, "8.4 สร้างรายงานจาก On-site")
    p(doc, "หลัง Leader Check-out แล้ว สามารถกดสร้าง Daily Report ให้สมาชิกใน session โดยใช้เวลาเริ่มและจบจากห้อง On-site เพื่อช่วยลดการกรอกรายงานซ้ำ")

    h1(doc, "9. Requests สำหรับ Manager/Admin")
    h2(doc, "9.1 อนุมัติคำขอ")
    numbered(doc, [
        "เปิด Requests",
        "ถ้าเป็น Manager/Admin จะเห็นแท็บ คำขอของฉัน และ รออนุมัติ",
        "เลือก OT Request หรือ Leave Request",
        "แตะรายการเพื่อดูรายละเอียด",
        "กดอนุมัติ หรือไม่อนุมัติพร้อมระบุเหตุผล",
    ])
    h2(doc, "9.2 ผลหลังอนุมัติใบลา")
    bullets(doc, [
        "ลาทั้งวันจะสร้างหรือปรับ daily_time_logs เป็น status = leave",
        "ลาครึ่งวันหรือรายชั่วโมงจะคำนวณ late threshold ใหม่ หากพนักงาน Check-in แล้ว",
        "การยกเลิกใบลาที่อนุมัติแล้วจากปุ่ม ยกเลิกใบลา จะเปลี่ยนเป็น cancelled ทันที คืน balance ทันที และเรียก DB RPC เพื่อล้าง/คำนวณ daily_time_logs ใหม่",
        "หากยังมีข้อมูลเก่าที่เป็น cancel_requested ค้างอยู่ ระบบยังรองรับให้ Manager/Admin จัดการต่อได้",
    ])

    h1(doc, "10. Calendar และ Profile")
    h2(doc, "10.1 Calendar")
    bullets(doc, [
        "ใช้ดูภาพรวมวันทำงาน วันลา วันหยุด และสถานะประจำวัน",
        "ช่วยให้พนักงานตรวจว่ามีวันใดขาดรายงานหรือผิดปกติ",
        "ใช้เทียบกับ Requests และ Daily Report ก่อนแจ้งแก้ไข",
    ])
    h2(doc, "10.2 Profile")
    bullets(doc, [
        "แสดงสรุปเดือนปัจจุบัน เช่น จำนวนวันมาสาย ชั่วโมง OT และอัตราส่งรายงาน",
        "มีปฏิทินส่งรายงานและประวัติรายวัน",
        "ใช้ตรวจ self-audit ก่อนสิ้นเดือน",
    ])

    h1(doc, "11. HR / Admin")
    h2(doc, "11.1 HR Attendance")
    bullets(doc, [
        "ดูสรุปพนักงานตามช่วงเดือน/วันที่เลือก",
        "เห็นสถานะมา, สาย, ขาด, ลา, วันหยุด และ OT",
        "คลิกพนักงานเพื่อดูรายละเอียดรายวัน",
        "รองรับ multi-select เพื่อ export สรุปรวมหรือ export รายละเอียดเป็นไฟล์ Excel",
        "แสดงเบี้ยเลี้ยง, ประเภทงาน, คนขับขาไป/ขากลับ และ OT จาก timeline/request",
    ])
    h2(doc, "11.2 Daily Audit")
    bullets(doc, [
        "ใช้ตรวจสอบและแก้ไขข้อมูลเวลาเฉพาะกรณี เช่น ลืม Check-in/Check-out หรือข้อมูลผิด",
        "การแก้ไขโดย Admin จะบันทึก timeline event เป็น admin override",
        "วันหยุดไม่นับสายเมื่อ Admin ปรับเวลา",
    ])
    h2(doc, "11.3 Time Sync")
    bullets(doc, [
        "ใช้ตรวจความครบถ้วนของ Daily Report เทียบกับเวลาทำงาน",
        "ช่วยหาเคสที่มี Check-in แต่ยังไม่ได้ส่งรายงาน",
        "เหมาะสำหรับปิดรอบงานรายวัน/รายเดือน",
    ])
    h2(doc, "11.4 Settings")
    table(
        doc,
        ["แท็บ", "หน้าที่"],
        [
            ("จัดการรายงาน", "ตั้งค่า End User, Project, รายละเอียดงานที่ใช้ใน Daily Report"),
            ("จัดการสิทธิ์", "เปลี่ยน role เป็น user/manager/admin และเปิด/พัก/ระงับสิทธิ์บัญชี"),
            ("วันหยุด", "จัดการวันหยุดและเสาร์ทำงาน เพื่อให้ระบบคำนวณวันและ OT ถูกต้อง"),
            ("นโยบายวันลา", "กำหนดประเภทลา สิทธิ์ และกฎพื้นฐาน"),
            ("วันลาพนักงาน", "ดูและแก้ไขยอดสิทธิ์ลา ใช้ไปแล้ว ยกยอด และสร้างสิทธิ์ปีปัจจุบัน"),
        ],
        widths=[4, 12],
    )

    h1(doc, "12. QR Display")
    bullets(doc, [
        "เป็นหน้าจอสำหรับแสดง QR Check-in โรงงาน",
        "ควรเปิดบนจอส่วนกลางหรืออุปกรณ์ที่บริษัทกำหนด",
        "QR จะ refresh อัตโนมัติและมีเวลาหมดอายุสั้น",
        "รายการคนที่สแกนล่าสุดจะแสดงบนหน้าจอเพื่อยืนยันเบื้องต้น",
        "เมนูนี้เป็น Admin-only และควรเปิดในอุปกรณ์ที่เชื่อถือได้",
    ])

    h1(doc, "13. แนวทางแก้ปัญหาที่พบบ่อย")
    table(
        doc,
        ["อาการ", "สาเหตุที่เป็นไปได้", "วิธีแก้"],
        [
            ("สแกน QR ไม่ได้", "QR หมดอายุ, กล้องไม่อนุญาต, QR ถูกใช้ซ้ำ", "รีเฟรช/รอ QR ใหม่ อนุญาตกล้อง และสแกนจากจอจริง"),
            ("ขึ้นว่า Check-in แล้ว", "มี first_check_in ของวันนี้แล้ว", "ตรวจหน้า Home/Profile หากผิดให้แจ้ง Admin ตรวจ Audit"),
            ("ปุ่ม OT ยังไม่เปิด", "ยังไม่ถึง 18:00 หรือสถานะงานยังไม่พร้อม", "รอถึงเวลาและตรวจว่า Check-out งานปกติแล้ว"),
            ("ขอใบลาแล้วสถานะยังไม่เปลี่ยน", "ยังรออนุมัติ", "ตรวจ Requests และรอ Manager/Admin action"),
            ("ลาพักร้อนขึ้นเตือน", "ยื่นล่วงหน้าน้อยกว่า 3 วัน", "ปรับวันที่หรือส่งให้หัวหน้าพิจารณาตามนโยบาย"),
            ("On-site ปิดไม่ได้", "สิทธิ์ไม่ใช่ Leader หรือ session ไม่อยู่สถานะที่ถูกต้อง", "ให้ Leader เข้าห้องและตรวจสถานะ session"),
            ("Daily Report บันทึกไม่ได้", "กรอก Detail/End User/Project/Period ไม่ครบ", "เติมช่องที่ยังว่างให้ครบทุกการ์ดงาน"),
            ("บัญชีเข้าใช้งานไม่ได้", "pending หรือ suspended", "ติดต่อ Admin ให้ตรวจ access_status"),
        ],
        widths=[4.5, 5.5, 6],
    )

    h1(doc, "14. Checklist สำหรับวันทำงานปกติ")
    table(
        doc,
        ["เวลา/ช่วง", "สิ่งที่ควรทำ", "หมายเหตุ"],
        [
            ("ก่อน 08:30", "สแกน QR Check-in", "หากมีใบลาช่วงเช้าให้ตรวจว่าใบลาอนุมัติแล้ว"),
            ("ระหว่างวัน", "ทำงานตามปกติ / บันทึก On-site หากออกนอกสถานที่", "ถ้าเปลี่ยนงานหลายโครงการให้จดไว้เพื่อทำ Daily Report"),
            ("ก่อนเลิกงาน", "กรอก Daily Report", "ควรกรอกในวันเดียวกัน"),
            ("17:30", "Check-out งานปกติ", "ถ้าลืม ระบบ Auto Check-out ให้ แต่ไม่ควรพึ่งระบบ"),
            ("หลัง 18:00", "Start OT เมื่อมี OT จริง", "เมื่อจบงานให้ End OT และยื่น OT Request หากจำเป็น"),
        ],
        widths=[3.5, 7, 5.5],
    )
    callout(doc, "สรุปสำหรับพนักงาน", "เช้า Check-in, เย็น Check-out, ทำงานนอกสถานที่ให้ใช้ On-site, มี OT/ลาให้ยื่น Requests, และส่ง Daily Report ให้ครบทุกวันที่ทำงาน", GREEN)

    h1(doc, "15. Checklist สำหรับผู้ดูแลการอบรม")
    bullets(doc, [
        "เตรียมบัญชีทดสอบอย่างน้อย 1 บัญชีสำหรับ Employee และ 1 บัญชีสำหรับ Manager/Admin",
        "เปิด QR Display ให้ผู้เข้าอบรมเห็น flow การสแกน",
        "สาธิตการยื่น OT และใบลา พร้อมการอนุมัติจาก Manager",
        "สาธิต On-site ตั้งแต่สร้างห้องจน Check-out และสร้างรายงาน",
        "สาธิต HR Attendance และ export รายงาน",
        "ย้ำกติกาเวลา 08:30, Auto Check-out 17:30, OT 18:00, และลาพักร้อนล่วงหน้า 3 วัน",
    ])

    add_footer(doc)
    doc.save(OUT)


if __name__ == "__main__":
    build_manual()
    print(OUT)
